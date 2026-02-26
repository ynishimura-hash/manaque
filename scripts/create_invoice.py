from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor
import datetime
import os
from dateutil.relativedelta import relativedelta

def create_invoice(date_str, amount_str, output_filename="請求書.docx"):
    doc = Document()
    
    # Title
    title = doc.add_heading('請 求 書', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph("") # Spacer

    # Top Section: Recipient and Sender
    # Use a table for layout
    table_bg = doc.add_table(rows=1, cols=2)
    table_bg.autofit = True
    
    # Recipient (Left)
    cell_to = table_bg.cell(0, 0)
    p_to = cell_to.paragraphs[0]
    p_to.add_run("合同会社EIS 御中").bold = True
    p_to.runs[0].font.size = Pt(16)
    cell_to.add_paragraph("件名：システム構築費")

    # Sender (Right)
    cell_from = table_bg.cell(0, 1)
    p_from = cell_from.paragraphs[0]
    p_from.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_from.add_run("西村 友祐").bold = True
    
    # Calculate Invoice Date (1 month before transfer date)
    transfer_date = datetime.datetime.strptime(date_str, "%Y/%m/%d")
    invoice_date = transfer_date - relativedelta(months=1)
    invoice_date_str = invoice_date.strftime("%Y年%m月%d日")
    
    date_p = cell_from.add_paragraph(f"請求日：{invoice_date_str}")
    date_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    doc.add_paragraph("") # Spacer
    doc.add_paragraph("") # Spacer

    # Grand Total
    total_p = doc.add_paragraph()
    total_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_total_label = total_p.add_run("ご請求金額　")
    run_total_label.font.size = Pt(14)
    run_total = total_p.add_run(f"¥ {amount_str} -")
    run_total.bold = True
    run_total.font.size = Pt(24)
    run_total.font.underline = True

    doc.add_paragraph("") # Spacer

    # Details Table
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '品目'
    hdr_cells[1].text = '数量'
    hdr_cells[2].text = '単価'
    hdr_cells[3].text = '金額'
    
    # Center align headers
    for cell in hdr_cells:
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    # Row 1
    row_cells = table.add_row().cells
    row_cells[0].text = 'システム構築費'
    row_cells[1].text = '1'
    row_cells[2].text = f"¥ {amount_str}"
    row_cells[3].text = f"¥ {amount_str}"

    # Add empty rows
    for _ in range(3):
        table.add_row()

    doc.add_paragraph("") # Spacer

    # Bank Info
    doc.add_heading('お振込先', level=2)
    bank_info = doc.add_paragraph()
    bank_info.add_run("ゆうちょ銀行\n").bold = True
    bank_info.add_run("六四八（ロクヨンハチ）店 (648)\n")
    bank_info.add_run("普通　1050996\n")
    bank_info.add_run("記号番号：16440　口座番号：10509961\n")
    bank_info.add_run("名義：ニシムラ ユウスケ")

    doc.add_paragraph("")
    note = doc.add_paragraph("※ 振込手数料は貴社負担にてお願いいたします。")
    note.runs[0].font.size = Pt(9)
    note.runs[0].font.color.rgb = RGBColor(100, 100, 100)

    # Save
    save_path = f"/Users/yuyu24/.gemini/antigravity/brain/4b9b954d-1a27-4fa2-b5a0-7dc54559babe/{output_filename}"
    doc.save(save_path)
    print(f"Generated: {save_path}")

if __name__ == "__main__":
    # Transactions extracted from PDF
    transactions = [
        ("2025/02/03", "300,000"),
        ("2025/04/25", "200,000"),
        ("2025/05/01", "260,000"),
        ("2025/05/21", "250,000"),
        ("2025/06/24", "250,000"),
        ("2025/07/25", "230,000"),
        ("2025/08/25", "200,000"),
        ("2025/11/26", "270,000"),
        ("2025/12/29", "150,000"),
    ]

    for date_str, amount in transactions:
        # Create filename like 請求書_20250203_300000.docx
        dt = datetime.datetime.strptime(date_str, "%Y/%m/%d")
        fname = f"請求書_{dt.strftime('%Y%m%d')}_{amount.replace(',', '')}.docx"
        create_invoice(date_str, amount, fname)
